#!/usr/bin/env python3
"""
模型评估报告生成器
扫描所有 module_eval 目录，为每个模型生成独立的 Markdown 和 Word 报告
"""

import json
import os
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Tuple

import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
import numpy as np
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 设置中文字体支持
matplotlib.rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial']
matplotlib.rcParams['axes.unicode_minus'] = False


class ModelEvaluator:
    """数据解析和评估指标计算"""

    def __init__(self, base_dir: str):
        self.base_dir = Path(base_dir)
        self.modules = []

    def scan_module_dirs(self) -> List[Path]:
        """扫描所有 module_eval 结尾的目录"""
        module_dirs = []
        for item in self.base_dir.iterdir():
            if item.is_dir() and item.name.endswith('_model_eval'):
                module_dirs.append(item)
        return sorted(module_dirs)

    @staticmethod
    def parse_json(module_dir: Path) -> Dict[str, Any]:
        """解析 test_metrics_summary_per_class.json"""
        json_path = module_dir / 'test_metrics_summary_per_class.json'
        if not json_path.exists():
            return None

        with open(json_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    @staticmethod
    def parse_csv(module_dir: Path) -> pd.DataFrame:
        """解析 test_metrics_detailed.csv"""
        csv_path = module_dir / 'test_metrics_detailed.csv'
        if not csv_path.exists():
            return None

        return pd.read_csv(csv_path)

    @staticmethod
    def get_plot_files(module_dir: Path) -> List[Path]:
        """获取所有 plot 图片文件"""
        plots_dir = module_dir / 'plots'
        if not plots_dir.exists():
            return []

        plot_files = sorted(plots_dir.glob('*.png'),
                           key=lambda x: int(re.search(r'(\d+)', x.stem).group()) if re.search(r'(\d+)', x.stem) else 0)
        return plot_files

    @staticmethod
    def extract_model_name(module_dir: Path) -> str:
        """从目录名提取可读的模型名称"""
        name = module_dir.name.replace('_model_eval', '')
        # 转换为可读格式
        name = name.replace('_', ' ').title()
        return name

    def load_module_data(self, module_dir: Path) -> Dict[str, Any]:
        """加载单个模块的所有数据"""
        return {
            'dir': module_dir,
            'name': self.extract_model_name(module_dir),
            'json_data': self.parse_json(module_dir),
            'csv_data': self.parse_csv(module_dir),
            'plots': self.get_plot_files(module_dir)
        }


class ChartGenerator:
    """统计图表生成器"""

    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_dice_boxplot(self, csv_data: pd.DataFrame, model_name: str) -> Path:
        """生成 Dice 系数箱线图"""
        fig, axes = plt.subplots(1, 3, figsize=(15, 5))

        dice_cols = ['dice_rv', 'dice_myo', 'dice_lv']
        labels = ['RV', 'Myocardium', 'LV']
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1']

        for i, (col, label, color) in enumerate(zip(dice_cols, labels, colors)):
            data = csv_data[col].dropna()
            axes[i].boxplot(data, vert=True, patch_artist=True,
                           boxprops=dict(facecolor=color, alpha=0.7),
                           medianprops=dict(color='red', linewidth=2))
            axes[i].set_title(f'{label}\nDice Coefficient Distribution', fontsize=12, fontweight='bold')
            axes[i].set_ylabel('Dice Score', fontsize=10)
            axes[i].grid(True, alpha=0.3)
            axes[i].set_ylim([0, 1.05])

            # 添加统计信息
            median_val = data.median()
            mean_val = data.mean()
            axes[i].text(0.5, 0.02, f'Median: {median_val:.3f}\nMean: {mean_val:.3f}',
                        transform=axes[i].transAxes, ha='center', va='bottom',
                        bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))

        plt.suptitle(f'{model_name} - Dice Score Distribution', fontsize=14, fontweight='bold')
        plt.tight_layout()

        output_path = self.output_dir / f'{model_name.replace(" ", "_").lower()}_dice_boxplot.png'
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()

        return output_path

    def generate_hd_boxplot(self, csv_data: pd.DataFrame, model_name: str) -> Path:
        """生成 Hausdorff 距离箱线图"""
        fig, axes = plt.subplots(1, 3, figsize=(15, 5))

        hd_cols = ['hd_rv', 'hd_myo', 'hd_lv']
        labels = ['RV', 'Myocardium', 'LV']
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1']

        for i, (col, label, color) in enumerate(zip(hd_cols, labels, colors)):
            data = csv_data[col].dropna()
            axes[i].boxplot(data, vert=True, patch_artist=True,
                           boxprops=dict(facecolor=color, alpha=0.7),
                           medianprops=dict(color='red', linewidth=2))
            axes[i].set_title(f'{label}\nHausdorff Distance Distribution', fontsize=12, fontweight='bold')
            axes[i].set_ylabel('Distance (pixels)', fontsize=10)
            axes[i].grid(True, alpha=0.3)

            # 添加统计信息
            median_val = data.median()
            mean_val = data.mean()
            axes[i].text(0.5, 0.02, f'Median: {median_val:.2f}\nMean: {mean_val:.2f}',
                        transform=axes[i].transAxes, ha='center', va='bottom',
                        bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))

        plt.suptitle(f'{model_name} - Hausdorff Distance Distribution', fontsize=14, fontweight='bold')
        plt.tight_layout()

        output_path = self.output_dir / f'{model_name.replace(" ", "_").lower()}_hd_boxplot.png'
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()

        return output_path

    def generate_radar_chart(self, json_data: Dict, model_name: str) -> Path:
        """生成综合指标雷达图"""
        categories = ['RV', 'Myocardium', 'LV']

        # 提取平均指标
        dice_means = [json_data['dice_mean'][cat] for cat in ['RV', 'Myocardium', 'LV']]
        f1_means = [json_data['metrics_mean'][cat]['f1_score'] for cat in ['RV', 'Myocardium', 'LV']]
        iou_means = [json_data['metrics_mean'][cat]['iou'] for cat in ['RV', 'Myocardium', 'LV']]
        accuracy_means = [json_data['metrics_mean'][cat]['accuracy'] for cat in ['RV', 'Myocardium', 'LV']]

        # 创建雷达图
        fig = plt.figure(figsize=(10, 10))
        ax = fig.add_subplot(111, projection='polar')

        angles = np.linspace(0, 2 * np.pi, len(categories), endpoint=False).tolist()
        angles += angles[:1]  # 闭合

        dice_means += dice_means[:1]
        f1_means += f1_means[:1]
        iou_means += iou_means[:1]
        accuracy_means += accuracy_means[:1]

        ax.plot(angles, dice_means, 'o-', linewidth=2, label='Dice', color='#FF6B6B')
        ax.fill(angles, dice_means, alpha=0.15, color='#FF6B6B')
        ax.plot(angles, f1_means, 'o-', linewidth=2, label='F1-Score', color='#4ECDC4')
        ax.fill(angles, f1_means, alpha=0.15, color='#4ECDC4')
        ax.plot(angles, iou_means, 'o-', linewidth=2, label='IoU', color='#45B7D1')
        ax.fill(angles, iou_means, alpha=0.15, color='#45B7D1')

        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(categories, fontsize=12)
        ax.set_ylim(0, 1)
        ax.set_yticks([0.2, 0.4, 0.6, 0.8, 1.0])
        ax.set_yticklabels(['0.2', '0.4', '0.6', '0.8', '1.0'], fontsize=9)
        ax.grid(True, linestyle='--', alpha=0.7)

        plt.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize=11)
        plt.title(f'{model_name}\nPerformance Metrics Overview',
                 size=14, fontweight='bold', pad=20)

        output_path = self.output_dir / f'{model_name.replace(" ", "_").lower()}_radar.png'
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()

        return output_path

    def generate_all_charts(self, module_data: Dict) -> Dict[str, Path]:
        """为单个模型生成所有图表"""
        model_name = module_data['name']
        csv_data = module_data['csv_data']
        json_data = module_data['json_data']

        charts = {}

        if csv_data is not None:
            charts['dice_boxplot'] = self.generate_dice_boxplot(csv_data, model_name)
            charts['hd_boxplot'] = self.generate_hd_boxplot(csv_data, model_name)

        if json_data is not None:
            charts['radar'] = self.generate_radar_chart(json_data, model_name)

        return charts


class MarkdownReporter:
    """Markdown 报告生成器"""

    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _generate_performance_description(self, json_data: Dict) -> str:
        """生成详细的性能描述文字"""
        categories = ['RV', 'Myocardium', 'LV']
        category_names = {
            'RV': '右心室 (RV)',
            'Myocardium': '心肌 (Myocardium)',
            'LV': '左心室 (LV)'
        }
        descriptions = []

        # Dice 排名
        dice_scores = {cat: json_data['dice_mean'][cat] for cat in categories}
        sorted_dice = sorted(dice_scores.items(), key=lambda x: x[1], reverse=True)
        rankings = {cat: i + 1 for i, (cat, _) in enumerate(sorted_dice)}

        for cat in categories:
            dice = json_data['dice_mean'][cat]
            dice_std = json_data['dice_std'][cat]
            hd = json_data['hd_mean'][cat]
            hd_std = json_data['hd_std'][cat]

            f1 = json_data['metrics_mean'][cat]['f1_score']
            iou = json_data['metrics_mean'][cat]['iou']
            recall = json_data['metrics_mean'][cat]['recall']

            # 稳定性评估
            if dice_std < 0.1:
                stability = "极佳的稳定性"
            elif dice_std < 0.2:
                stability = "良好的稳定性"
            elif dice_std < 0.3:
                stability = "中等稳定性"
            else:
                stability = "较低稳定性"

            # 表现等级
            if dice >= 0.9:
                performance = "卓越"
            elif dice >= 0.8:
                performance = "优秀"
            elif dice >= 0.7:
                performance = "良好"
            elif dice >= 0.6:
                performance = "中等"
            else:
                performance = "需要改进"

            ranking = rankings[cat]
            ranking_zh = {1: "一", 2: "二", 3: "三"}[ranking]

            desc = f"""
**{category_names[cat]} 分割：**
- Dice 系数：**{dice:.4f}**（标准差：{dice_std:.4f}），表现**{performance}**，具有**{stability}**
- 在三个类别中基于 Dice 得分排名：第 **{ranking_zh}** 名
- Hausdorff 距离：**{hd:.2f} ± {hd_std:.2f}** 像素
- F1-Score：**{f1:.4f}**，IoU：**{iou:.4f}**，Recall：**{recall:.4f}**
"""
            descriptions.append(desc)

        return '\n'.join(descriptions)

    def _generate_metrics_table(self, json_data: Dict) -> str:
        """生成性能指标表格"""
        categories = ['RV', 'Myocardium', 'LV']
        category_names = {
            'RV': '右心室 (RV)',
            'Myocardium': '心肌 (Myocardium)',
            'LV': '左心室 (LV)'
        }

        table = "| 类别 | Dice (均值±标准差) | HD (均值±标准差) | 准确率 | 召回率 | 特异性 | F1-Score | IoU |\n"
        table += "|------|------------------|-----------------|--------|--------|--------|----------|-----|\n"

        for cat in categories:
            dice_mean = json_data['dice_mean'][cat]
            dice_std = json_data['dice_std'][cat]
            hd_mean = json_data['hd_mean'][cat]
            hd_std = json_data['hd_std'][cat]

            acc = json_data['metrics_mean'][cat]['accuracy']
            rec = json_data['metrics_mean'][cat]['recall']
            spec = json_data['metrics_mean'][cat]['specificity']
            f1 = json_data['metrics_mean'][cat]['f1_score']
            iou = json_data['metrics_mean'][cat]['iou']

            table += f"| {category_names[cat]} | {dice_mean:.4f} ± {dice_std:.4f} | {hd_mean:.2f} ± {hd_std:.2f} | {acc:.4f} | {rec:.4f} | {spec:.4f} | {f1:.4f} | {iou:.4f} |\n"

        return table

    def generate_report(self, module_data: Dict, charts: Dict[str, Path]) -> Path:
        """生成单个模型的 Markdown 报告"""
        model_name = module_data['name']
        json_data = module_data['json_data']
        plots = module_data['plots']

        # 使用文件系统安全的名称
        safe_name = module_data['dir'].name.replace('_model_eval', '')
        output_path = self.output_dir / f'{safe_name}.md'

        with open(output_path, 'w', encoding='utf-8') as f:
            # 标题
            f.write(f"# {model_name} - 评估报告\n\n")
            f.write(f"**生成时间：** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write("---\n\n")

            # 基本信息
            f.write("## 📋 模型信息\n\n")
            f.write(f"- **模型类型：** {json_data['model_type']}\n")
            f.write(f"- **总参数量：** {json_data['total_parameters']:,}\n")
            f.write(f"- **测试样本数：** {json_data['num_test_samples']}\n\n")

            # 性能指标表格
            f.write("## 📊 性能指标\n\n")
            f.write(self._generate_metrics_table(json_data))
            f.write("\n")

            # 详细性能描述
            f.write("## 📈 性能分析\n\n")
            f.write(self._generate_performance_description(json_data))
            f.write("\n")

            # 统计图表
            f.write("## 📉 统计图表\n\n")

            if 'dice_boxplot' in charts:
                f.write("### Dice 得分分布\n\n")
                rel_path = charts['dice_boxplot'].relative_to(self.output_dir.parent.parent)
                f.write(f"![Dice 箱线图](./{rel_path})\n\n")

            if 'hd_boxplot' in charts:
                f.write("### Hausdorff 距离分布\n\n")
                rel_path = charts['hd_boxplot'].relative_to(self.output_dir.parent.parent)
                f.write(f"![HD 箱线图](./{rel_path})\n\n")

            if 'radar' in charts:
                f.write("### 综合性能概览\n\n")
                rel_path = charts['radar'].relative_to(self.output_dir.parent.parent)
                f.write(f"![雷达图](./{rel_path})\n\n")

            # 可视化图片
            if plots:
                f.write("## 🖼️ 样本可视化\n\n")
                for plot_path in plots:
                    rel_path = plot_path.relative_to(self.output_dir.parent.parent)
                    f.write(f"![{plot_path.stem}](./{rel_path})\n\n")

        return output_path


class WordReporter:
    """Word 报告生成器"""

    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _add_model_info(self, doc: Document, json_data: Dict):
        """添加模型基本信息"""
        doc.add_heading('模型信息', level=2)

        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Light Grid Accent 1'

        info_table.rows[0].cells[0].text = '模型类型'
        info_table.rows[0].cells[1].text = json_data['model_type']
        info_table.rows[1].cells[0].text = '总参数量'
        info_table.rows[1].cells[1].text = f"{json_data['total_parameters']:,}"
        info_table.rows[2].cells[0].text = '测试样本数'
        info_table.rows[2].cells[1].text = str(json_data['num_test_samples'])

        doc.add_paragraph()

    def _add_metrics_table(self, doc: Document, json_data: Dict):
        """添加性能指标表格"""
        doc.add_heading('性能指标', level=2)

        table = doc.add_table(rows=4, cols=8)
        table.style = 'Light Grid Accent 1'

        # 表头
        headers = ['类别', 'Dice (均值±标准差)', 'HD (均值±标准差)', '准确率',
                   '召回率', '特异性', 'F1-Score', 'IoU']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # 设置表头样式
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

        # 数据行
        categories = ['RV', 'Myocardium', 'LV']
        category_names = {
            'RV': '右心室 (RV)',
            'Myocardium': '心肌 (Myocardium)',
            'LV': '左心室 (LV)'
        }
        for i, cat in enumerate(categories):
            row = i + 1
            dice_mean = json_data['dice_mean'][cat]
            dice_std = json_data['dice_std'][cat]
            hd_mean = json_data['hd_mean'][cat]
            hd_std = json_data['hd_std'][cat]

            table.rows[row].cells[0].text = category_names[cat]
            table.rows[row].cells[1].text = f"{dice_mean:.4f} ± {dice_std:.4f}"
            table.rows[row].cells[2].text = f"{hd_mean:.2f} ± {hd_std:.2f}"
            table.rows[row].cells[3].text = f"{json_data['metrics_mean'][cat]['accuracy']:.4f}"
            table.rows[row].cells[4].text = f"{json_data['metrics_mean'][cat]['recall']:.4f}"
            table.rows[row].cells[5].text = f"{json_data['metrics_mean'][cat]['specificity']:.4f}"
            table.rows[row].cells[6].text = f"{json_data['metrics_mean'][cat]['f1_score']:.4f}"
            table.rows[row].cells[7].text = f"{json_data['metrics_mean'][cat]['iou']:.4f}"

        doc.add_paragraph()

    def _add_performance_analysis(self, doc: Document, json_data: Dict):
        """添加详细性能分析"""
        doc.add_heading('性能分析', level=2)

        categories = ['RV', 'Myocardium', 'LV']
        category_names = {
            'RV': '右心室 (RV)',
            'Myocardium': '心肌 (Myocardium)',
            'LV': '左心室 (LV)'
        }
        ranking_zh_map = {1: '一', 2: '二', 3: '三'}

        # Dice 排名
        dice_scores = {cat: json_data['dice_mean'][cat] for cat in categories}
        sorted_dice = sorted(dice_scores.items(), key=lambda x: x[1], reverse=True)
        rankings = {cat: i + 1 for i, (cat, _) in enumerate(sorted_dice)}

        for cat in categories:
            dice = json_data['dice_mean'][cat]
            dice_std = json_data['dice_std'][cat]
            hd = json_data['hd_mean'][cat]
            hd_std = json_data['hd_std'][cat]

            f1 = json_data['metrics_mean'][cat]['f1_score']
            iou = json_data['metrics_mean'][cat]['iou']
            recall = json_data['metrics_mean'][cat]['recall']

            # 稳定性评估
            if dice_std < 0.1:
                stability = "极佳的稳定性"
            elif dice_std < 0.2:
                stability = "良好的稳定性"
            elif dice_std < 0.3:
                stability = "中等稳定性"
            else:
                stability = "较低稳定性"

            # 表现等级
            if dice >= 0.9:
                performance = "卓越"
            elif dice >= 0.8:
                performance = "优秀"
            elif dice >= 0.7:
                performance = "良好"
            elif dice >= 0.6:
                performance = "中等"
            else:
                performance = "需要改进"

            ranking = rankings[cat]
            ranking_zh = ranking_zh_map[ranking]

            p = doc.add_paragraph()
            p.add_run(f'{category_names[cat]} 分割：\n').bold = True
            p.add_run(
                f'Dice 系数为 {dice:.4f}（标准差：{dice_std:.4f}），'
                f'表现{performance}，具有{stability}。\n'
            )
            p.add_run(f'在三个类别中基于 Dice 得分排名：第 {ranking_zh} 名。\n')
            p.add_run(
                f'Hausdorff 距离为 {hd:.2f} ± {hd_std:.2f} 像素。'
                f'F1-Score 为 {f1:.4f}，IoU 为 {iou:.4f}，Recall 为 {recall:.4f}'
            )

            doc.add_paragraph()  # 空行

    def _add_charts(self, doc: Document, charts: Dict[str, Path]):
        """添加统计图表"""
        doc.add_heading('统计图表', level=2)

        chart_labels = {
            'dice_boxplot': 'Dice 得分分布',
            'hd_boxplot': 'Hausdorff 距离分布',
            'radar': '综合性能概览'
        }

        for key, label in chart_labels.items():
            if key in charts:
                doc.add_heading(label, level=3)
                doc.add_picture(str(charts[key]), width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

    def _add_visualizations(self, doc: Document, plots: List[Path]):
        """添加可视化图片（2x5 网格）"""
        if not plots:
            return

        doc.add_heading('样本可视化', level=2)
        doc.add_paragraph('以下展示了分割结果的样本图片：\n')

        # 创建 2x5 表格
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Light Grid Accent 1'

        for i, plot_path in enumerate(plots):
            row = i // 5
            col = i % 5
            cell = table.rows[row].cells[col]

            # 清除单元格内容
            cell.text = ''
            # 添加图片
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(str(plot_path), width=Inches(1.8))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加标题
            title_para = cell.add_paragraph(f'样本 {i}')
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.runs[0]
            run.font.size = Pt(9)

    def generate_report(self, module_data: Dict, charts: Dict[str, Path]) -> Path:
        """生成单个模型的 Word 报告"""
        model_name = module_data['name']
        json_data = module_data['json_data']
        plots = module_data['plots']

        # 使用文件系统安全的名称
        safe_name = module_data['dir'].name.replace('_model_eval', '')
        output_path = self.output_dir / f'{safe_name}.docx'

        doc = Document()

        # 设置默认字体
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')

        # 标题
        title = doc.add_heading(f'{model_name} - 评估报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}',
            style='Intense Quote'
        )
        doc.add_paragraph()

        # 添加各个部分
        self._add_model_info(doc, json_data)
        self._add_metrics_table(doc, json_data)
        self._add_performance_analysis(doc, json_data)
        self._add_charts(doc, charts)
        self._add_visualizations(doc, plots)

        doc.save(str(output_path))
        return output_path


def main():
    """主流程"""
    # 设置路径
    base_dir = Path('/home/charliechen/final_project/baseline/result')
    reports_dir = base_dir / 'reports'
    charts_dir = reports_dir / 'charts'

    print("=" * 60)
    print("🚀 Starting Model Evaluation Report Generation")
    print("=" * 60)

    # 初始化各个模块
    evaluator = ModelEvaluator(str(base_dir))
    chart_gen = ChartGenerator(charts_dir)
    md_reporter = MarkdownReporter(reports_dir)
    word_reporter = WordReporter(reports_dir)

    # 扫描模块目录
    module_dirs = evaluator.scan_module_dirs()
    print(f"\n📁 Found {len(module_dirs)} model directories:")
    for d in module_dirs:
        print(f"   - {d.name}")

    print(f"\n📊 Reports will be saved to: {reports_dir}")
    print(f"📈 Charts will be saved to: {charts_dir}\n")

    # 为每个模型生成报告
    for i, module_dir in enumerate(module_dirs, 1):
        print(f"\n{'=' * 60}")
        print(f"📝 Processing [{i}/{len(module_dirs)}]: {module_dir.name}")
        print(f"{'=' * 60}")

        # 加载数据
        module_data = evaluator.load_module_data(module_dir)

        if module_data['json_data'] is None:
            print(f"⚠️  Skipping {module_dir.name}: No JSON data found")
            continue

        print(f"   Model: {module_data['name']}")
        print(f"   Plots found: {len(module_data['plots'])}")

        # 生成图表
        print("   📈 Generating charts...")
        charts = chart_gen.generate_all_charts(module_data)
        print(f"      Charts created: {list(charts.keys())}")

        # 生成 Markdown 报告
        print("   📄 Generating Markdown report...")
        md_path = md_reporter.generate_report(module_data, charts)
        print(f"      ✅ {md_path.name}")

        # 生成 Word 报告
        print("   📝 Generating Word report...")
        word_path = word_reporter.generate_report(module_data, charts)
        print(f"      ✅ {word_path.name}")

    print(f"\n{'=' * 60}")
    print("✨ All reports generated successfully!")
    print(f"{'=' * 60}")
    print(f"\n📂 Output directory: {reports_dir}")
    print(f"\nGenerated files:")

    # 列出生成的文件
    for f in sorted(reports_dir.iterdir()):
        if f.is_file():
            size = f.stat().st_size / 1024  # KB
            print(f"   - {f.name} ({size:.1f} KB)")


if __name__ == '__main__':
    main()
